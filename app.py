import os
import io
import csv
import time
import streamlit as st
from dotenv import load_dotenv
from datetime import datetime
from typing import List

import pandas as pd
import PyPDF2
from docx import Document as DocxDocument

# LangChain / embeddings / vectorstore imports (attempt)
EMBEDDINGS_OK = False
PINECONE_OK = False
HUGGINGFACE_EMBEDDINGS = None
PINECONE = None

try:
    from langchain_community.embeddings import HuggingFaceEmbeddings
    HUGGINGFACE_EMBEDDINGS = HuggingFaceEmbeddings
    EMBEDDINGS_OK = True
except Exception:
    EMBEDDINGS_OK = False

try:
    import pinecone
    from langchain_community.vectorstores import Pinecone as PineconeStore
    pinecone_api_key = os.getenv("PINECONE_API_KEY")
    pinecone_env = os.getenv("PINECONE_ENVIRONMENT") or "us-west1-gcp"
    if pinecone_api_key:
        pinecone.init(api_key=pinecone_api_key, environment=pinecone_env)
        PINECONE = PineconeStore
        PINECONE_OK = True
    else:
        PINECONE_OK = False
except Exception:
    PINECONE_OK = False

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_groq import ChatGroq
from langchain.schema import Document
from langchain.memory import ConversationBufferWindowMemory
import re

load_dotenv()
GROQ_API_KEY = os.getenv("GROQ_API_KEY")

PRIMARY = "#0b93f6"

HEADER_LOGO_URL = "https://cdn-icons-png.flaticon.com/512/4712/4712027.png"
ASSISTANT_LOGO_URL = "https://cdn-icons-png.flaticon.com/512/4712/4712027.png"

st.set_page_config(page_title="SAP Ariba RAG Chatbot", layout="wide")
st.markdown(f"""
    <style>
    .header-card {{ background-color: {PRIMARY}; padding:16px; border-radius:10px; color: white; display:flex; gap:16px; align-items:center; }}
    .chat-user {{
        background: linear-gradient(90deg, rgba(11,147,246,0.12), rgba(11,147,246,0.08));
        border-left: 4px solid {PRIMARY};
        padding:10px; border-radius:10px; margin:8px 0; display:flex; gap:8px; align-items:flex-start;
        color: #fff;
    }}
    .chat-assistant {{
        background:#f3f4f6; padding:10px; border-radius:10px; margin:8px 0; display:flex; gap:8px; align-items:flex-start;
        color: #222;
    }}
    .icon-left {{ width:28px; height:28px; margin-top:2px; }}
    .quiz-card {{ border:1px solid #eee; padding:14px; border-radius:10px; margin-bottom:12px; background: #fff; box-shadow: 0 2px 6px rgba(0,0,0,0.03); }}
    .question-badge {{ display:inline-block; background:{PRIMARY}; color:white; padding:6px 10px; border-radius:12px; font-weight:700; margin-right:10px; }}
    .correct {{ color: green; font-weight:700; }}
    .summary-card {{ border-left: 6px solid #0b93f6; padding: 12px; border-radius: 8px; background: #2d2d2d; color: #ffffff; box-shadow: 0 2px 6px rgba(0,0,0,0.5); margin-bottom: 12px; line-height: 1.6; }}
    .small-muted {{ color:#666; font-size:13px; }}
    .active-file {{ font-weight:700; color: {PRIMARY}; }}
    .chat-options-bar {{
        display: flex;
        align-items: center;
        background: #181818;
        border-radius: 8px;
        padding: 8px 12px;
        margin-bottom: 10px;
        gap: 8px;
        box-shadow: 0 2px 8px rgba(0,0,0,0.07);
    }}
    .chat-options-bar .option-select {{
        background: #232323;
        color: #fff;
        border-radius: 6px;
        border: none;
        padding: 6px 10px;
        font-size: 15px;
        margin-right: 6px;
    }}
    .chat-options-bar .attach-btn {{
        background: #232323;
        color: #fff;
        border: none;
        border-radius: 6px;
        padding: 8px 10px;
        font-size: 15px;
        cursor: pointer;
        margin-right: 6px;
    }}
    .chat-options-bar .send-btn {{
        background: #0b93f6;
        color: #fff;
        border: none;
        border-radius: 6px;
        padding: 8px 16px;
        font-size: 15px;
        cursor: pointer;
        transition: background 0.2s;
        margin-right: 6px;
    }}
    .chat-options-bar .clear-btn {{
        background: #444;
        color: #fff;
        border: none;
        border-radius: 6px;
        padding: 8px 16px;
        font-size: 15px;
        cursor: pointer;
    }}
    .mic-button {{
        background: #ff4757;
        color: #fff;
        border: none;
        border-radius: 50%;
        width: 40px;
        height: 40px;
        font-size: 18px;
        cursor: pointer;
        transition: all 0.3s ease;
        display: flex;
        align-items: center;
        justify-content: center;
    }}
    .mic-button:hover {{
        background: #ff3838;
        transform: scale(1.1);
    }}
    .mic-button.recording {{
        background: #ff1744;
        animation: pulse 1.5s infinite;
    }}
    @keyframes pulse {{
        0% {{ box-shadow: 0 0 0 0 rgba(255, 23, 68, 0.7); }}
        70% {{ box-shadow: 0 0 0 10px rgba(255, 23, 68, 0); }}
        100% {{ box-shadow: 0 0 0 0 rgba(255, 23, 68, 0); }}
    }}
    .audio-controls {{
        display: flex;
        gap: 8px;
        align-items: center;
        margin-bottom: 10px;
    }}
    .speaker-button {{
        background: #2ed573;
        color: #fff;
        border: none;
        border-radius: 6px;
        padding: 6px 12px;
        font-size: 14px;
        cursor: pointer;
        transition: all 0.2s;
    }}
    .speaker-button:hover {{
        background: #26d05f;
        transform: translateY(-1px);
    }}
    .audio-status {{
        font-size: 12px;
        color: #666;
        margin-left: 10px;
    }}
    </style>
""", unsafe_allow_html=True)

# JavaScript for Web Speech API and Speech Synthesis
st.markdown("""
<script>
let recognition = null;
let speechSynthesis = window.speechSynthesis;
let isRecording = false;

// Initialize Speech Recognition
if ('webkitSpeechRecognition' in window || 'SpeechRecognition' in window) {
    const SpeechRecognition = window.SpeechRecognition || window.webkitSpeechRecognition;
    recognition = new SpeechRecognition();
    recognition.continuous = false;
    recognition.interimResults = false;
    recognition.lang = 'en-US';
    
    recognition.onresult = function(event) {
        const transcript = event.results[0][0].transcript;
        const inputElement = parent.document.querySelector('[data-testid="stTextInput"] input');
        if (inputElement) {
            inputElement.value = transcript;
            inputElement.dispatchEvent(new Event('input', { bubbles: true }));
            inputElement.dispatchEvent(new Event('change', { bubbles: true }));
        }
        stopRecording();
    };
    
    recognition.onerror = function(event) {
        console.error('Speech recognition error:', event.error);
        stopRecording();
    };
    
    recognition.onend = function() {
        stopRecording();
    };
}

function startRecording() {
    if (recognition && !isRecording) {
        isRecording = true;
        recognition.start();
        updateMicButton(true);
    }
}

function stopRecording() {
    if (recognition && isRecording) {
        isRecording = false;
        recognition.stop();
        updateMicButton(false);
    }
}

function updateMicButton(recording) {
    const micButton = parent.document.querySelector('.mic-button');
    if (micButton) {
        if (recording) {
            micButton.classList.add('recording');
            micButton.innerHTML = '⏹️';
        } else {
            micButton.classList.remove('recording');
            micButton.innerHTML = '🎤';
        }
    }
}

function speakText(text) {
    if (speechSynthesis) {
        // Stop any ongoing speech
        speechSynthesis.cancel();
        
        const utterance = new SpeechSynthesisUtterance(text);
        utterance.rate = 0.9;
        utterance.pitch = 1;
        utterance.volume = 0.8;
        
        // Try to use a female voice if available
        const voices = speechSynthesis.getVoices();
        const femaleVoice = voices.find(voice => 
            voice.name.toLowerCase().includes('female') || 
            voice.name.toLowerCase().includes('zira') ||
            voice.name.toLowerCase().includes('eva') ||
            voice.gender === 'female'
        );
        
        if (femaleVoice) {
            utterance.voice = femaleVoice;
        }
        
        utterance.onstart = function() {
            updateSpeakerStatus('Speaking...');
        };
        
        utterance.onend = function() {
            updateSpeakerStatus('Ready');
        };
        
        utterance.onerror = function() {
            updateSpeakerStatus('Error');
        };
        
        speechSynthesis.speak(utterance);
    }
}

function stopSpeaking() {
    if (speechSynthesis) {
        speechSynthesis.cancel();
        updateSpeakerStatus('Stopped');
    }
}

function updateSpeakerStatus(status) {
    const statusElement = parent.document.querySelector('.audio-status');
    if (statusElement) {
        statusElement.textContent = status;
    }
}

// Make functions available globally
window.startRecording = startRecording;
window.stopRecording = stopRecording;
window.speakText = speakText;
window.stopSpeaking = stopSpeaking;
</script>
""", unsafe_allow_html=True)

QNA_SYSTEM = """
You are SAP Ariba Expert Assistant. Use ONLY the provided context passages.
- If the topic or keyword appears in the context, return the relevant details from the context.
- Do NOT invent facts or use outside knowledge.
- If the answer is not present in the provided context, reply exactly: "I could not find the answer in the provided material."
- When you reference a fact, add a friendly source note like: (from filename.pdf)
Answer in a simple, customer-friendly tone.
"""

SUMMARY_SYSTEM = """
You are SAP Ariba Expert Assistant. Using ONLY the provided context, OUTPUT in Markdown:

### Summary
<3–6 sentences overview>

### Study Notes
- 5–7 concise bullets

### Key Definitions
- **Term**: short definition

If insufficient info, say: "The provided material does not contain enough information to summarize fully."
"""

QUIZ_SYSTEM = """
You are SAP Ariba Expert Assistant. Using ONLY the provided context, create EXACTLY 5 multiple-choice questions.
Output strictly in this format for each question:

Q1. Question text
A. option
B. option
C. option
D. option
Answer: X

Rules:
- Do not use outside knowledge.
- Make questions clear and relevant to the context.
"""

def load_csv_safely(uploaded_file) -> pd.DataFrame | None:
    try:
        sample = uploaded_file.read(200000).decode("utf-8", errors="ignore")
        uploaded_file.seek(0)
        dialect = csv.Sniffer().sniff(sample, delimiters=[",", ";", "\t", "|"])
        sep = dialect.delimiter
    except Exception:
        sep = ","
    try:
        df = pd.read_csv(uploaded_file, sep=sep, on_bad_lines="skip", engine="python", quotechar='"')
        return df
    except Exception:
        try:
            uploaded_file.seek(0)
            df = pd.read_csv(uploaded_file, on_bad_lines="skip", engine="python")
            return df
        except Exception:
            return None

def extract_text_from_pdf(uploaded_file) -> str:
    try:
        reader = PyPDF2.PdfReader(uploaded_file)
        text = ""
        for i, page in enumerate(reader.pages):
            try:
                page_text = page.extract_text() or ""
                text += f"\n--- Page {i+1} ---\n{page_text}\n"
            except Exception:
                continue
        return text
    except Exception:
        return ""

def extract_text_from_docx(uploaded_file) -> str:
    try:
        doc = DocxDocument(uploaded_file)
        out = []
        for p in doc.paragraphs:
            if p.text and p.text.strip():
                out.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells])
                if row_text:
                    out.append(row_text)
        return "\n".join(out)
    except Exception:
        return ""

from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
class TFIDFWrapper:
    def __init__(self):
        self.texts = []
        self.metadatas = []
        self.vectorizer = None
        self.matrix = None

    def add_documents(self, documents: List[Document]):
        for d in documents:
            self.texts.append(d.page_content)
            self.metadatas.append(d.metadata if d.metadata else {})
        self.vectorizer = TfidfVectorizer(stop_words='english', max_features=5000)
        self.matrix = self.vectorizer.fit_transform(self.texts)

    def from_documents(self, documents: List[Document]):
        self.texts = [d.page_content for d in documents]
        self.metadatas = [d.metadata if d.metadata else {} for d in documents]
        self.vectorizer = TfidfVectorizer(stop_words='english', max_features=5000)
        self.matrix = self.vectorizer.fit_transform(self.texts)

    def get_relevant_documents(self, query: str, k: int = 3) -> List[Document]:
        if self.matrix is None or self.vectorizer is None or len(self.texts) == 0:
            return []
        q_vec = self.vectorizer.transform([query])
        sims = cosine_similarity(q_vec, self.matrix)[0]
        top_idxs = sims.argsort()[::-1][:k]
        results = []
        for i in top_idxs:
            results.append(Document(page_content=self.texts[i], metadata=self.metadatas[i]))
        return results

class EnhancedRAGChatbot:
    def __init__(self):
        self.embeddings = None
        self.vectorstore = None
        self.llm = None
        self.chain = None
        self.memory = None
        self.conversation_history = []
        self._init_components()

    def _init_components(self):
        if EMBEDDINGS_OK and PINECONE_OK:
            try:
                self.embeddings = HUGGINGFACE_EMBEDDINGS(
                    model_name="sentence-transformers/all-MiniLM-L6-v2",
                    model_kwargs={"device": "cpu"}
                )
            except Exception:
                self.embeddings = None
        else:
            
            self.embeddings = None

        groq_key = os.getenv("GROQ_API_KEY") or GROQ_API_KEY
        if groq_key:
            try:
                model_name = st.session_state.get("selected_model", "gemma2-9b-it")
                temp = st.session_state.get("temperature", 0.1)
                self.llm = ChatGroq(groq_api_key=groq_key, model_name=model_name, temperature=temp)
            except Exception:
                self.llm = None
        else:
            self.llm = None

        try:
            self.memory = ConversationBufferWindowMemory(k=5, return_messages=True)
        except Exception:
            self.memory = None

        if self.llm:
            prompt_template = ChatPromptTemplate.from_messages([
                ("system", QNA_SYSTEM + """

CONTEXT INFORMATION:
{context}

CONVERSATION HISTORY:
{chat_history}

Current Date: {current_date}"""),
                ("user", "{question}")
            ])
            output_parser = StrOutputParser()
            try:
                self.chain = prompt_template | self.llm | output_parser
            except Exception:
                self.chain = None

        self._load_default_documents()

    def _load_default_documents(self):
        docs = [
            Document(page_content="SAP Ariba Contract Management Process details...", metadata={"source": "Contract_Management_Guide"}),
            Document(page_content="SAP Ariba Sourcing Process details...", metadata={"source": "Sourcing_Process_Guide"}),
        ]
        if self.embeddings is not None and PINECONE_OK:
            try:
                index_name = "ariba-chatbot"
                if index_name not in pinecone.list_indexes():
                    pinecone.create_index(index_name, dimension=384)
                self.vectorstore = PINECONE.from_documents(docs, self.embeddings, index_name=index_name)
            except Exception:
                self.vectorstore = TFIDFWrapper()
                self.vectorstore.from_documents(docs)
        else:
            self.vectorstore = TFIDFWrapper()
            self.vectorstore.from_documents(docs)

    def add_documents(self, documents: List[Document]):
        if self.vectorstore is None:
            self.vectorstore = TFIDFWrapper()
        try:
            if PINECONE_OK and hasattr(self.vectorstore, "add_documents") and not isinstance(self.vectorstore, TFIDFWrapper):
                self.vectorstore.add_documents(documents)
            elif isinstance(self.vectorstore, TFIDFWrapper):
                self.vectorstore.add_documents(documents)
            else:
                self.vectorstore = TFIDFWrapper()
                self.vectorstore.from_documents(documents)
        except Exception:
            self.vectorstore = TFIDFWrapper()
            self.vectorstore.from_documents(documents)

    def chat(self, question: str) -> str:
        if self.vectorstore is None:
            return "No documents indexed. Upload documents first."
            try:
                if PINECONE_OK and hasattr(self.vectorstore, "as_retriever") and not isinstance(self.vectorstore, TFIDFWrapper):
                    retriever = self.vectorstore.as_retriever(search_kwargs={"k": 3})
                    docs = retriever.get_relevant_documents(question)
                elif isinstance(self.vectorstore, TFIDFWrapper):
                    docs = self.vectorstore.get_relevant_documents(question, k=3)
                else:
                    if hasattr(self.vectorstore, "get_relevant_documents"):
                        docs = self.vectorstore.get_relevant_documents(question)
                    else:
                        docs = []
            except Exception as e:
                return f"Retriever error: {e}"
                
            max_context_chars = 4000
            context_text = "\n\n".join([d.page_content for d in docs])[:max_context_chars]
            chat_history = "\n".join(self.conversation_history[-10:])
            current_date = datetime.now().strftime("%Y-%m-%d")
            
            if self.chain is None:
                if context_text.strip():
                    answer = f"Context found from documents:\n\n{context_text[:1000]}...\n\n(LLM unavailable — set GROQ_API_KEY to enable LLM answers.)"
                else:
                    answer = "No contextual documents found and LLM not available."
            else:
                try:
                    answer = self.chain.invoke({
                        "context": context_text,
                        "chat_history": chat_history,
                        "question": question,
                        "current_date": current_date
                    })
                except Exception as e:
                    answer = f"LLM/chain error: {e}"
            self.conversation_history.append(f"User: {question}")
            self.conversation_history.append(f"AI: {answer}")
            return answer

# Initialize session state
if "chatbot" not in st.session_state:
    st.session_state.chatbot = EnhancedRAGChatbot()
if "messages" not in st.session_state:
    st.session_state.messages = []
if "user_input" not in st.session_state:
    st.session_state.user_input = ""
if "selected_model" not in st.session_state:
    st.session_state.selected_model = "gemma2-9b-it"
if "temperature" not in st.session_state:
    st.session_state.temperature = 0.1

# Initialize Summary & Quiz Feature State
if "summary_output" not in st.session_state:
    st.session_state.summary_output = None
if "quiz_questions" not in st.session_state:
    st.session_state.quiz_questions = None
if "quiz_index" not in st.session_state:
    st.session_state.quiz_index = 0
if "quiz_score" not in st.session_state:
    st.session_state.quiz_score = 0
if "quiz_done" not in st.session_state:
    st.session_state.quiz_done = False
if "quiz_feedback" not in st.session_state:
    st.session_state.quiz_feedback = None
if "is_recording" not in st.session_state:
    st.session_state.is_recording = False
if "audio_enabled" not in st.session_state:
    st.session_state.audio_enabled = False

def process_uploaded_files(uploaded_files):
    docs = []
    for f in uploaded_files:
        name = f.name.lower()
        content = ""
        if name.endswith(".pdf"):
            content = extract_text_from_pdf(f)
        elif name.endswith(".docx"):
            content = extract_text_from_docx(f)
        elif name.endswith(".txt"):
            try:
                f.seek(0)
                content = f.read().decode("utf-8", errors="ignore")
            except Exception:
                content = ""
        elif name.endswith(".csv"):
            f.seek(0)
            df = load_csv_safely(f)
            if df is not None:
                content = df.to_string(index=False)
        elif name.endswith((".xls", ".xlsx")):
            try:
                f.seek(0)
                df = pd.read_excel(f, engine="openpyxl")
                content = df.to_string(index=False)
            except Exception:
                content = ""
        else:
            content = ""
        if content and content.strip():
            docs.append(Document(page_content=content, metadata={"source": f.name}))

    if docs:
        st.session_state.chatbot.add_documents(docs)
        st.success(f"Indexed {len(docs)} uploaded files.")
    else:
        st.warning("No content extracted from uploaded files.")

def stream_assistant_text(text: str, placeholder: st.delta_generator.DeltaGenerator):
    """Stream the assistant's response below the latest user message, like ChatGPT."""
    words = text.split()
    out = ""
    for w in words:
        out += w + " "
        # Stream Markdown content inside your styled chat bubble
        placeholder.markdown(
            f"""
            <div class='chat-assistant'>
                <div class='icon-left'>
                    <img src='{ASSISTANT_LOGO_URL}' class='icon-left'/>
                </div>
                <div>
                    <strong>SAP Ariba Chatbot:</strong>
                    <div style="margin-top: 5px;">
                        {st.markdown(out)}
                    </div>
                </div>
            </div>
            """,
            unsafe_allow_html=True
        )
        time.sleep(0.02)

    # Final render
    placeholder.markdown(
        f"""
        <div class='chat-assistant'>
            <div class='icon-left'>
                <img src='{ASSISTANT_LOGO_URL}' class='icon-left'/>
            </div>
            <div>
                <strong>SAP Ariba Chatbot:</strong>
                <div style="margin-top: 5px;">
                    {st.markdown(out)}
                </div>
            </div>
        </div>
        """,
        unsafe_allow_html=True
    )

def on_send():
    text = st.session_state.user_input.strip()
    if not text:
        st.warning("Please enter a message.")
        return

    # Save user message
    st.session_state.messages.append({
        "role": "user",
        "content": text,
        "timestamp": datetime.now().strftime("%H:%M:%S")
    })

    try:
        # Create LLM for selected mode
        llm = ChatGroq(
            groq_api_key=GROQ_API_KEY,
            model_name=st.session_state.get("selected_model", "gemma2-9b-it"),
            temperature=st.session_state.get("temperature", 0.7 if st.session_state.chat_mode == "General" else 0.1)
        )
    except Exception:
        llm = None

    with st.spinner("Generating answer..."):
        if st.session_state.chat_mode == "General":
            # General ChatGPT-like mode
            if llm:
                if "general_memory" not in st.session_state:
                    st.session_state.general_memory = ConversationBufferWindowMemory(k=10, return_messages=True)

                prompt_template = ChatPromptTemplate.from_messages([
                    ("system", "You are a helpful, knowledgeable, and friendly assistant. Answer naturally and conversationally."),
                    ("user", "{question}")
                ])
                chain = prompt_template | llm | StrOutputParser()
                try:
                    resp = chain.invoke({"question": text})
                except Exception as e:
                    resp = f"General mode error: {e}"
            else:
                resp = "LLM not available — please set GROQ_API_KEY."
        
        else:
            # Original RAG mode
            try:
                st.session_state.chatbot.llm = llm
                prompt_template = ChatPromptTemplate.from_messages([
                    ("system", QNA_SYSTEM + """

CONTEXT INFORMATION:
{context}

CONVERSATION HISTORY:
{chat_history}

Current Date: {current_date}"""),
                    ("user", "{question}")
                ])
                st.session_state.chatbot.chain = prompt_template | st.session_state.chatbot.llm | StrOutputParser()
            except Exception:
                pass

            resp = st.session_state.chatbot.chat(text)

    # Save assistant message
    st.session_state.messages.append({
        "role": "assistant",
        "content": resp,
        "timestamp": datetime.now().strftime("%H:%M:%S")
    })

    # Trigger streaming effect
    st.session_state.show_streaming = True

    # Auto-speak if enabled
    if st.session_state.get("audio_enabled", False):
        clean_response = re.sub(r'\*\*(.+?)\*\*', r'\1', resp)
        clean_response = re.sub(r'\*(.+?)\*', r'\1', clean_response)
        clean_response = re.sub(r'[#\-\*\[\]()]', '', clean_response)
        st.session_state.speak_text = clean_response

    # Reset input
    st.session_state.user_input = ""

def toggle_mic():
    st.session_state.is_recording = not st.session_state.get("is_recording", False)

def toggle_audio():
    st.session_state.audio_enabled = not st.session_state.get("audio_enabled", False)

def on_clear():
    st.session_state.messages = []
    st.session_state.chatbot.conversation_history = []

with st.sidebar:
    st.image(HEADER_LOGO_URL, width=80)
    st.title("SAP Ariba RAG Chatbot")
    uploaded_files = st.file_uploader("Upload documents (pdf/docx/txt/csv/xlsx)", accept_multiple_files=True, type=["pdf","docx","txt","csv","xls","xlsx"])
    if uploaded_files:
        if st.button("Process & Index Uploaded Files"):
            process_uploaded_files(uploaded_files)
    st.markdown("---")
    st.subheader("LLM Settings")
    st.session_state.temperature = st.slider("Temperature", 0.0, 1.0, st.session_state.temperature, key="temperature_slider")
    st.session_state.selected_model = st.selectbox("Model", ["gemma2-9b-it", "mixtral-8x7b-32768", "llama3-8b-8192"], index=0, key="model_select")
    st.markdown("---")
    st.subheader("Diagnostics")
    if EMBEDDINGS_OK:
        st.success("Embeddings (HuggingFace) available")
    else:
        st.warning("Embeddings not available — using TF-IDF fallback")
    if PINECONE_OK:
        st.success("Pinecone available")
    else:
        st.info("Pinecone not available — using TF-IDF fallback store")
    st.markdown("<div class='footer'>Made with ❤️ using Streamlit & LangChain</div>", unsafe_allow_html=True)

st.markdown(f"""
    <div class="header-card">
        <div style="flex-shrink:0;">
            <img src="{HEADER_LOGO_URL}" class="icon-left"/>
        </div>
        <div>
            <h1 style="color:white;margin:0;">SAP Ariba RAG Chatbot</h1>
            <p style="color:#f0f0f0;margin:4px 0 0 0;">Your SAP Ariba expert assistant — upload documents, ask questions, summarize, and quiz yourself.</p>
        </div>
    </div>
""", unsafe_allow_html=True)

import re

def markdown_to_html(text):
    text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)
    text = re.sub(r'\*(.+?)\*', r'<i>\1</i>', text)
    return text

# Summary & Quiz UI Panel - Added before chat interface
with st.expander("📄 Summary & Quiz Tools", expanded=False):
    col1, col2 = st.columns(2)
    with col1:
        st.markdown("#### Generate Summary")

        # Summary mode selection
        summary_mode = st.radio(
            "Summary mode",
            ["Summarize whole uploaded documents", "Summarize a specific topic"],
            index=0,
            key="summary_mode"
        )

        # Always show typing option
        summary_topic = st.text_input("Enter instructions or topic (optional):", key="summary_topic")

        summary_btn = st.button("Generate Summary", key="summary_btn")
        if summary_btn:
            query = summary_topic.strip() if summary_topic else "summarize documents"    
            if hasattr(st.session_state.chatbot.vectorstore, "get_relevant_documents"):
                docs = st.session_state.chatbot.vectorstore.get_relevant_documents(query, k=6)
            else: 
                docs = []
            context_text = "\n\n".join([d.page_content for d in docs])[:4000]

            if st.session_state.chatbot.llm:
                prompt_template = ChatPromptTemplate.from_messages([
                    ("system", SUMMARY_SYSTEM),
                    ("user", "{question}")
                ])
                chain = prompt_template | st.session_state.chatbot.llm | StrOutputParser()
                try:
                    summary = chain.invoke({
                        "question": f"Summarize {'the uploaded documents' if summary_mode=='Summarize whole uploaded documents' else 'the specific topic'} {('with instructions: ' + summary_topic) if summary_topic else ''}.\n\nContext:\n{context_text}"
                    })
                except Exception as e:
                    summary = f"Summary error: {e}"
            else:
                summary = "LLM not available for summary."
            st.session_state.summary_output = summary

    with col2:
        st.markdown("#### Generate Quiz")

        # Quiz mode selection
        quiz_mode = st.radio(
            "Quiz mode",
            ["Create quiz from whole documents", "Create quiz from a specific topic"],
            index=0,
            key="quiz_mode"
        )

        # Always show typing option
        quiz_topic = st.text_input("Enter quiz topic or instructions (optional):", key="quiz_topic")

        quiz_btn = st.button("Generate Quiz", key="quiz_btn")
        if quiz_btn:
            query = quiz_topic.strip() if quiz_topic else "quiz"
            if hasattr(st.session_state.chatbot.vectorstore, "get_relevant_documents"):
                docs = st.session_state.chatbot.vectorstore.get_relevant_documents(query, k=6)
            else:
                docs = []
            context_text = "\n\n".join([d.page_content for d in docs])[:4000]

            if st.session_state.chatbot.llm:
                prompt_template = ChatPromptTemplate.from_messages([
                    ("system", QUIZ_SYSTEM),
                    ("user", "{question}")
                ])
                chain = prompt_template | st.session_state.chatbot.llm | StrOutputParser()
                try:
                    quiz_raw = chain.invoke({
                        "question": f"Create a quiz {('from the uploaded documents' if quiz_mode=='Create quiz from whole documents' else 'on the specific topic')} {('with instructions: ' + quiz_topic) if quiz_topic else ''}.\n\nContext:\n{context_text}"
                    })
                except Exception as e:
                    quiz_raw = f"Quiz error: {e}"
            else:
                quiz_raw = "LLM not available for quiz."

            # Parse quiz questions
            import re
            quiz_qs = []
            if isinstance(quiz_raw, str):
                pattern = r"Q\d+\.(.*?)\nA\.(.*?)\nB\.(.*?)\nC\.(.*?)\nD\.(.*?)\nAnswer:\s*([A-D])"
                matches = re.findall(pattern, quiz_raw, re.DOTALL)
                for i, m in enumerate(matches[:5]):
                    import re
                    quiz_qs.append({
                        "question": re.sub(r"\*+", "", m[0].strip()),
                        "options": [m[1].strip(), m[2].strip(), m[3].strip(), m[4].strip()],
                        "answer_index": "ABCD".index(m[5].strip())
                    })
            st.session_state.quiz_questions = quiz_qs
            st.session_state.quiz_index = 0
            st.session_state.quiz_score = 0
            st.session_state.quiz_done = False
            st.session_state.quiz_feedback = None

# Summary Display UI - Added after expander, before quiz

if st.session_state.summary_output:
    st.markdown("---")
    st.header("📄 Summary")
    
    # Create a container with proper styling for the summary
    summary_container = st.container()
    with summary_container:
        # Display summary content with proper styling
        st.markdown("""
        <div style='background-color: #262730; padding: 20px; border-radius: 8px; border-left: 4px solid #0ea5e9; margin: 10px 0; color: white;'>
            <div style='color: #888; font-size: 14px; margin-bottom: 15px;'>
                The provided material does not contain enough information to summarize fully.
            </div>
        </div>
        """, unsafe_allow_html=True)
        
        # Display the actual summary content using st.markdown without HTML wrapper
        with st.container():
            st.markdown(st.session_state.summary_output)
        
        # Add a button to clear the summary with better layout
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            if st.button("Clear Summary", key="clear_summary", use_container_width=True):
                st.session_state.summary_output = None
                st.rerun()

# Quiz Runner UI - Added before chat interface
if st.session_state.quiz_questions:
    questions = st.session_state.quiz_questions
    idx = st.session_state.quiz_index
    total = len(questions)
    st.markdown("---")
    st.header("📝 Quiz")
    if st.session_state.quiz_done:
        st.success(f"Quiz completed! Your score: {st.session_state.quiz_score} / {total}")
        if st.button("Restart Quiz"):
            st.session_state.quiz_index = 0
            st.session_state.quiz_score = 0
            st.session_state.quiz_done = False
            st.session_state.quiz_feedback = None
            st.rerun()
    else:
        q = questions[idx]
        st.markdown(f"**Question {idx+1} of {total}:**")
        st.write(q['question'])
        choice_key = f"quiz_choice_{idx}"
        selected = st.radio("Select an option:", q['options'], key=choice_key)
        submit_key = f"submit_{idx}"
        if st.button("Submit Answer", key=submit_key):
            sel_index = q['options'].index(selected)
            correct_index = int(q.get('answer_index', 0))
            if sel_index == correct_index:
                st.session_state.quiz_feedback = {"correct": True, "message": "Correct ✅"}
                st.session_state.quiz_score += 1
            else:
                st.session_state.quiz_feedback = {
                    "correct": False,
                    "message": f"Wrong ❌  | Correct: Option {correct_index+1}: {q['options'][correct_index]}"
                }
            st.rerun()
        fb = st.session_state.get('quiz_feedback')
        if fb:
            if fb.get('correct'):
                st.success(fb.get('message'))
            else:
                st.error(fb.get('message'))
            if idx + 1 < total:
                if st.button("Next Question"):
                    st.session_state.quiz_index += 1
                    st.session_state.quiz_feedback = None
                    st.rerun()
            else:
                if st.button("Finish Quiz"):
                    st.session_state.quiz_done = True
                    st.rerun()

# Audio Controls
with st.container():
    st.markdown('<div class="audio-controls">', unsafe_allow_html=True)
    audio_col1, audio_col2, audio_col3, audio_col4 = st.columns([1, 1, 2, 1])
    
    with audio_col1:
        audio_enabled = st.checkbox("🔊 Auto-speak", value=st.session_state.get("audio_enabled", False), key="audio_toggle")
        st.session_state.audio_enabled = audio_enabled
    
    with audio_col2:
        if st.button("🔇 Stop", key="stop_speaking"):
            st.markdown('<script>window.stopSpeaking && window.stopSpeaking();</script>', unsafe_allow_html=True)
    
    with audio_col3:
        st.markdown('<span class="audio-status">Ready</span>', unsafe_allow_html=True)
    
    st.markdown('</div>', unsafe_allow_html=True)

# ----------- Modern Chat Options Bar with Model, Mode, Mic, Attach, Send -----------
with st.container():
    chat_cols = st.columns([1.5, 1.5, 0.8, 0.8, 3.5, 0.7, 0.7])
    with chat_cols[0]:
        st.selectbox(
            "Model",
            ["gemma2-9b-it", "mixtral-8x7b-32768", "llama3-8b-8192"],
            index=["gemma2-9b-it", "mixtral-8x7b-32768", "llama3-8b-8192"].index(st.session_state.selected_model) if "selected_model" in st.session_state else 0,
            key="selected_model",
            label_visibility="collapsed",
            format_func=lambda x: f"🧠 {x}"
        )
    with chat_cols[1]:
        st.selectbox(
            "Mode",
            ["Ask", "Chat", "Search", "General"],  # ✅ Added General mode
            index=["Ask", "Chat", "Search", "General"].index(st.session_state.get("chat_mode", "Ask")),
            key="chat_mode",
            label_visibility="collapsed",
            format_func=lambda x: f"💬 {x}"
        )
    with chat_cols[2]:
        mic_button_class = "mic-button recording" if st.session_state.get("is_recording", False) else "mic-button"
        mic_icon = "⏹️" if st.session_state.get("is_recording", False) else "🎤"
        if st.button(mic_icon, key="mic_btn", help="Click to start/stop voice input"):
            st.markdown('''
            <script>
            if (window.startRecording && window.stopRecording) {
                if (!window.isRecording) {
                    window.startRecording();
                } else {
                    window.stopRecording();
                }
            } else {
                alert('Speech recognition not supported in this browser. Please use Chrome, Edge, or Safari.');
            }
            </script>
            ''', unsafe_allow_html=True)
    with chat_cols[3]:
        attach_clicked = st.button("📎", key="attach_btn", use_container_width=True)
        if attach_clicked:
            st.session_state.show_attach = True
    with chat_cols[4]:
        st.text_input(
            "User Input",
            key="user_input",
            placeholder="Type your question or click 🎤 to speak...",
            label_visibility="collapsed"
        )
    with chat_cols[5]:
        st.button("➤", on_click=on_send, use_container_width=True, key="send_btn")
    with chat_cols[6]:
        st.button("⨉", on_click=on_clear, use_container_width=True, key="clear_btn")

if st.session_state.get("show_attach", False):
    st.file_uploader("Attach file", type=["pdf","docx","txt","csv","xls","xlsx"], key="chat_attach", accept_multiple_files=True)
    if st.session_state.get("chat_attach"):
        process_uploaded_files(st.session_state.chat_attach)
        st.session_state.show_attach = False

# Display chat messages
if st.session_state.messages:
    for i, message in enumerate(st.session_state.messages):
        content = message["content"] or ""
        content_html = markdown_to_html(content)
        safe_html = content_html.replace("\n", "<br>")
        
        if message["role"] == "user":
            st.markdown(
                f"<div class='chat-user'><div class='icon-left'>🧑‍💼</div><div><strong>You:</strong> {safe_html}</div></div>", 
                unsafe_allow_html=True
            )
        elif message["role"] == "assistant":
            # For the last assistant message, show streaming effect
            if i == len(st.session_state.messages) - 1 and st.session_state.get("show_streaming", False):
                placeholder = st.empty()
                stream_assistant_text(content, placeholder)
                st.session_state.show_streaming = False
            else:
                # Add speak button for each assistant message
                col1, col2 = st.columns([0.95, 0.05])
                with col1:
                    st.markdown(
                        f"<div class='chat-assistant'><div class='icon-left'><img src='{ASSISTANT_LOGO_URL}' class='icon-left'/></div><div><strong>SAP Ariba Chatbot:</strong> {safe_html}</div></div>",
                        unsafe_allow_html=True
                    )
                with col2:
                    if st.button("🔊", key=f"speak_{i}", help="Click to hear this response"):
                        clean_text = re.sub(r'\*\*(.+?)\*\*', r'\1', content)
                        clean_text = re.sub(r'\*(.+?)\*', r'\1', clean_text)
                        clean_text = re.sub(r'[#\-\*\[\]()]', '', clean_text)
                        st.markdown(f'''
                        <script>
                        if (window.speakText) {{
                            window.speakText(`{clean_text.replace("`", "").replace("'", "").replace('"', '')}`);
                        }}
                        </script>
                        ''', unsafe_allow_html=True)

# Auto-speak the latest response if enabled
if st.session_state.get("speak_text") and st.session_state.get("audio_enabled", False):
    st.markdown(f'''
    <script>
    if (window.speakText) {{
        setTimeout(() => {{
            window.speakText(`{st.session_state.speak_text.replace("`", "").replace("'", "").replace('"', '')}`);
        }}, 1000);
    }}
    </script>
    ''', unsafe_allow_html=True)
    del st.session_state.speak_text

